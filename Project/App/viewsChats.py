from django.shortcuts import render, HttpResponse, redirect
from django.contrib.auth.decorators import login_required # Realizando isso ele só irá abrir a minha página se eu estiver logado.
from django.contrib.auth import authenticate, login, logout
#from tensorflow.keras.utils import to_categorical
#from tensorflow.keras.preprocessing.text import Tokenizer
#from tensorflow.keras.preprocessing.sequence import pad_sequences
#from tensorflow.keras.models import Sequential
#from tensorflow.keras.layers import Embedding, LSTM, Dense
from django.core.mail import send_mail, EmailMessage
from django.contrib import messages
from App.models import Evento
from django.conf import settings
from datetime import datetime
from docx import Document
from django.conf import settings
from spellchecker import SpellChecker # Biblioteca corrige acentuação da frase
import numpy as np
import logging
import uuid
import os
import re
logger = logging.getLogger('django')


## Bloco ChatBot cliente ------------------------------------------------------------------------------------------------------------------------------------------------------
def chatbot(request):
    return render(request, 'chatbot.html')

# Função para ler o conteúdo do documento Word
def ler_documento_word(caminhos_arquivos):
    texto = []
    for caminho_arquivo in caminhos_arquivos:
        doc = Document(caminho_arquivo)
        for par in doc.paragraphs:
            texto.append(par.text)
    texto_completo = '\n'.join(texto)
    return texto_completo


def buscar_resposta(texto, pergunta):
    # Verifica se a pergunta e o texto não são nulos
    if not texto or not pergunta:
        return "Pergunta ou documento inválidos."

    # Palavras-chave específicas para buscar no texto
    palavras_chave = [
        "comnect", "conexões de rede", "terminais", "aplicativos", "wireless", 
        "networks", "ws time out", "ws", "wireless networks",
        "time-out", "time", "out", "assembly", "cartao", "invalido",
        "time-out server", "cnpj", "windows installer",
        "instalar", "tls", "instalar o tls", "vpn slin", 
        "vpn slin não conecta", "tef", "criador", "criado", 
        "vpn", "portal erros", "erros portal", "erro portal"
    ]

    # Converte a pergunta e o texto para minúsculas para uma busca mais robusta
    pergunta = pergunta.lower()
    texto = texto.lower()
    start = None

    # Processa palavras-chave e faz busca no texto
    for palavra in palavras_chave:
        padrao = r'\b' + re.escape(palavra) + r'\b'  # Combinação exata
        if re.search(padrao, pergunta):
            match = re.search(padrao, texto)
            if match:
                start = match.start()
            if start is not None and start != -1:
                # Busca o fim do primeiro parágrafo após a palavra-chave
                end_paragrafo = texto.find('\n\n', start)
                # Se não encontrar a marcação de fim de parágrafo, limita a 5000 caracteres
                if end_paragrafo == -1:
                    end_paragrafo = start + 5000
                # Obtém o trecho correspondente ao primeiro parágrafo
                trecho = texto[start:end_paragrafo].strip()
                # Garante que o trecho comece com a primeira letra maiúscula
                trecho = trecho.capitalize() if trecho else trecho
                # Adiciona um ponto final se não houver pontuação ao final
                if not trecho.endswith(('.', '!', '?')):
                    trecho += '.'

                # Capitaliza a primeira letra após um ponto final
                trecho = re.sub(r'(?<=[.!?–°(])\s*(\w)', lambda x: x.group(0).upper(), trecho)
                return trecho

    return "Desculpe, não consegui encontrar uma resposta. Para mais informações, abra um chamado com o nosso suporte."

def chatbot_view(request):
    resposta = None  # Inicializa com None para ser usada no template
    if request.method == 'POST':
        pergunta = request.POST.get('pergunta')
        pergunta = pergunta.lower() if pergunta else ''
        
        # Se a pergunta for uma saudação
        if pergunta == "ola tudo bem?":
            resposta = "Eu estou ótimo! Obrigado por perguntar. \U0001F64F"
        else:
            if pergunta:
                caminhos_arquivos = [
                    os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Sobre a Comnect.docx'),
                    os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Lista de Erros.docx'),
                    os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Procedimentos Comnect.docx'),
                    os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Dúvidas.docx'),
                ]
                texto_documento = ler_documento_word(caminhos_arquivos)
                resposta = buscar_resposta(texto_documento, pergunta) if texto_documento else "Desculpe, não consegui encontrar uma resposta."

    return render(request, 'chatbot.html', {'resposta': resposta})


## Bloco ChatBot Interno ------------------------------------------------------------------------------------------------------------------------------------------------------

def chatbotInterno(request):
    return render(request, 'chatbotInterno.html')

def extrair_texto(caminho_arquivo):
    """Extrai o texto de um documento Word."""
    texto_completo = []
    try:
        doc = Document(caminho_arquivo)
        for par in doc.paragraphs:
            if par.text.strip():  # Ignorar parágrafos vazios
                texto_completo.append(par.text)
    except Exception as e:
        print(f"Erro ao processar o arquivo {caminho_arquivo}: {e}")
    return "\n".join(texto_completo)

def buscar_respostaInterno(texto, pergunta):
    """Busca uma resposta no texto com base em palavras-chave na pergunta."""
    if not texto or not pergunta:
        return None

    palavras_chave = ["Simulador PDV", "simulador", "teste de bancada", "Bancada", "Teste", "Queda", "Link", "Queda link", "queda link", "tunel ipsec", "tunel ipsec adquirente",
                      "comnect", "conexões de rede", "terminais", "aplicativos", "wirelles", "networks", "Ws time out", "ws", "wirelles networks",
                      "time-out", "time", "out", "assembly", "cartao", "invalido", "time-out server", "cnpj", "windows installer", "windows",
                      "invalido", "instalar", "tls", "instalar o tls", "vpn slin", "vpn slin não conecta", "criador", "criado", "vpn",
                      "portal erros", "erros portal", "erro portal", "cardse", "alteracao cadastro cliente no portal", "cadastro de loja cacau show", "loja cacau show", 
                      "sniffer fortgate", "carga de tabelas", "bandeiras", "verificacao de bandeiras", "relatorios no portal", "portal", "como realizar sniffer fortgate",
                      "extrair relatorio sitef", "relatorio sitef", "relatorio sitef web", "cadastro adquirente cacau show"
                     ]

    pergunta = pergunta.lower()
    texto = texto.lower()
    start = None

    for palavra in palavras_chave:
        padrao = r'\b' + re.escape(palavra) + r'\b'  # Combinação exata
        if re.search(padrao, pergunta):
            match = re.search(padrao, texto)
            if match:
                start = match.start()
                end_paragrafo = texto.find('\n\n', start)
                if end_paragrafo == -1:
                    end_paragrafo = start + 5000
                trecho = texto[start:end_paragrafo].strip().capitalize()
                if not trecho.endswith(('.', '!', '?')):
                    trecho += '.'
                trecho = re.sub(r'(?<=[.!?–°\n-])\s*(\w)', lambda x: x.group(0).upper(), trecho)
                return trecho

    return None

def chatbotInterno_view(request):
    resposta = None
    link_arquivo = None

    if request.method == 'POST':
        pergunta = request.POST.get('pergunta')
        if pergunta:
            caminhos_arquivos = [
                # Caminhos base de dados POPs
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP001 - Cadastro de Nova Loja - Cacau Show.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP002 - Teste de Bancada.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP004 - Verificação queda de Link.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP005 - Verificação de túnel IPsec de adquirente.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP006 - Cadastro de contrato CardSE - Cacau Show.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP008 - Como realizar um sniffer de pacotes no Fortigate.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP009 - Aplicação de Carga de Tabelas e Visualizador de Tabelas - Cacau Show.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP011 - Verificação de bandeiras.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP013 - Extrair relatório do SiTef via Sitef Web.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/DocsInterno', 'POP014 - Cadastro de Adquirentes - Cacau Show.docx'),
                # Caminhos perguntas default
                os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Sobre a Comnect.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Lista de Erros.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Procedimentos Comnect.docx'),
                os.path.join(settings.BASE_DIR, 'DataBase/Docs', 'Dúvidas.docx'),
                
            ]

            # Iterar sobre os arquivos
            for caminho_arquivo in caminhos_arquivos:
                if os.path.exists(caminho_arquivo):
                    # Extrair o texto do arquivo atual
                    texto_documento = extrair_texto(caminho_arquivo)
                    
                    # Buscar a resposta no texto do arquivo atual
                    resposta = buscar_respostaInterno(texto_documento, pergunta)
                    
                    # Se uma resposta for encontrada, configurar o link e interromper a busca
                    if resposta:
                        nome_arquivo = os.path.basename(caminho_arquivo)
                        link_arquivo = os.path.join(settings.MEDIA_URL, nome_arquivo)
                        break

    return render(request, 'chatbotInterno.html', {'resposta': resposta, 'link_arquivo': link_arquivo})

## Bloco chat --------------------------------------------------------------------------------------------------------------------------------------------------------------

def chat(request):
    return render(request, 'chat.html')